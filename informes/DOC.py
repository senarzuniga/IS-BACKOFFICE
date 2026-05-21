# =============================================================================
# INGECART — EXECUTIVE INDUSTRIAL REPORT GENERATOR
# Generates a PROFESSIONAL WORD REPORT (.DOCX)
# =============================================================================
#
# REQUIREMENTS:
# pip install python-docx pillow
#
# OUTPUT:
# Professional executive report with:
#
# ✓ Industrial premium styling
# ✓ INGECART branding
# ✓ Corporate orange highlights
# ✓ Header logo
# ✓ Footer with Ing_RESEARCH + page numbers
# ✓ Professional chapter styling
# ✓ Executive typography
# ✓ Tables with industrial aesthetics
#
# =============================================================================

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =============================================================================
# CONFIG
# =============================================================================

LOGO_PATH = r"C:\Users\Inaki Senar\Documents\INGECART\MARKETING\LOGOS\ICON ORANGE.png"

OUTPUT_DOCX = "INGECART_TradeShows_Executive_Report.docx"

# =============================================================================
# CORPORATE COLORS
# =============================================================================

BLACK = RGBColor(5, 7, 11)
ORANGE = RGBColor(255, 106, 0)
WHITE = RGBColor(244, 245, 247)
GREY = RGBColor(126, 132, 142)
DARK_GREY = RGBColor(26, 29, 36)

# =============================================================================
# DOCUMENT
# =============================================================================

doc = Document()

section = doc.sections[0]

section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

# =============================================================================
# HEADER
# =============================================================================

header = section.header

header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

try:
    run = header_para.add_run()
    run.add_picture(LOGO_PATH, width=Inches(0.9))
except:
    print("WARNING: Logo file not found")

# =============================================================================
# FOOTER
# =============================================================================

footer = section.footer

table = footer.add_table(rows=1, cols=2, width=Inches(6))
table.autofit = True

# LEFT FOOTER
left = table.cell(0, 0)
p_left = left.paragraphs[0]

run = p_left.add_run("Ing_RESEARCH")
run.font.name = "Montserrat"
run.font.size = Pt(9)
run.font.color.rgb = GREY

# RIGHT FOOTER PAGE NUMBER
right = table.cell(0, 1)
p_right = right.paragraphs[0]
p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

run = p_right.add_run()

fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.text = "PAGE"

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

# =============================================================================
# COVER PAGE
# =============================================================================

doc.add_paragraph("")

title = doc.add_paragraph()

run = title.add_run(
    "ESTUDIO EXHAUSTIVO DE FERIAS Y TRADE SHOWS"
)

run.font.name = "Montserrat"
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = BLACK

subtitle = doc.add_paragraph()

run = subtitle.add_run(
    "Sector Cartón Corrugado y Conversion | Enfoque Ingecart"
)

run.font.name = "Inter"
run.font.size = Pt(14)
run.font.color.rgb = ORANGE

doc.add_paragraph("")

date_p = doc.add_paragraph()

run = date_p.add_run("Fecha de elaboración: 2026-05-21")

run.font.name = "Inter"
run.font.size = Pt(11)
run.font.color.rgb = GREY

# =============================================================================
# DIVIDER LINE
# =============================================================================

divider = doc.add_paragraph()

pPr = divider._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')

bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '18')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'FF6A00')

pBdr.append(bottom)
pPr.append(pBdr)

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def add_chapter(title_text):

    p = doc.add_paragraph()

    run = p.add_run(title_text)

    run.font.name = "Montserrat"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = BLACK

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '14')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'FF6A00')

    pBdr.append(bottom)
    pPr.append(pBdr)

def add_subheading(text):

    p = doc.add_paragraph()

    run = p.add_run(text)

    run.font.name = "Montserrat"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = DARK_GREY

def add_body(text):

    p = doc.add_paragraph()

    run = p.add_run(text)

    run.font.name = "Inter"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(10)

def add_bullets(items):

    for item in items:

        p = doc.add_paragraph(style='List Bullet')

        run = p.add_run(item)

        run.font.name = "Inter"
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK

# =============================================================================
# CONTENT
# =============================================================================

add_chapter("1. OBJETIVO Y ALCANCE")

add_body(
    "Este estudio consolida y amplifica el listado base del archivo "
    "eventos_SHOWS_FERIAS_2025.xlsx con investigación sectorial adicional "
    "para los mercados estratégicos definidos por Ingecart."
)

add_subheading("Mercados analizados")

add_bullets([
    "Europa",
    "España",
    "USA",
    "México",
    "Canadá",
    "Australia",
    "South Korea"
])

add_subheading("Elementos incluidos")

add_bullets([
    "Fechas confirmadas o aproximadas",
    "Foco de cada evento",
    "Perfil de expositores",
    "Perfil de visitantes",
    "Fuentes y enlaces sectoriales",
    "Priorización estratégica para Ingecart"
])

# =============================================================================

add_chapter("2. CONTEXTO ESTRATÉGICO")

add_body(
    "El estudio ha sido priorizado utilizando el playbook comercial y técnico "
    "interno de Ingecart, orientado a integración de procesos, reducción de "
    "scrap, estabilización operativa y mejora de OEE."
)

add_subheading("Implicaciones estratégicas")

add_bullets([
    "Prioridad a eventos con decisores industriales",
    "Foco en corrugado, converting y automatización",
    "Alta relevancia para operaciones con cuellos de botella",
    "Preferencia por eventos con integración industrial real"
])

# =============================================================================

add_chapter("3. FUENTES EXPERTAS UTILIZADAS")

add_subheading("Publicaciones y asociaciones")

sources = [
    "The Packaging Portal",
    "Board Converting News",
    "FEFCO",
    "Corrugated of Course",
    "AICC",
    "TAPPI Corrugated Week",
    "SuperCorrExpo",
    "CCCA Canada"
]

add_bullets(sources)

# =============================================================================

add_chapter("4. EVENTOS DE ALTA RELEVANCIA PARA INGECART")

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

headers = ["EVENTO", "REGIÓN", "SCORE", "ENCaje ESTRATÉGICO"]

hdr_cells = table.rows[0].cells

for i, h in enumerate(headers):

    p = hdr_cells[i].paragraphs[0]

    run = p.add_run(h)

    run.font.name = "Montserrat"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = WHITE

    tcPr = hdr_cells[i]._tc.get_or_add_tcPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1A1D24')

    tcPr.append(shd)

events = [
    ["Corrugated Expo Barcelona", "Europa", "92", "Máximo fit corrugado"],
    ["FEFCO Summit Madrid", "Europa", "89", "Acceso C-Level"],
    ["Corrugated Week USA", "USA", "88", "Alta densidad industrial"],
    ["Empack Madrid", "España", "84", "Pipeline Iberia"],
    ["ALLFORPACK Paris", "Europa", "81", "Posicionamiento internacional"]
]

for event in events:

    row = table.add_row().cells

    for i, val in enumerate(event):

        p = row[i].paragraphs[0]

        run = p.add_run(val)

        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK

# =============================================================================

add_chapter("5. PLAN DE ACCIÓN 12 MESES")

add_subheading("Fase 1 — Captura de demanda")

add_bullets([
    "Activar mensajes verticales orientados a OEE y estabilidad",
    "Crear auditoría de flujo paquetizada",
    "Objetivo de 25-40 reuniones pre-feria"
])

add_subheading("Fase 2 — Conversión comercial")

add_bullets([
    "Pipeline estandarizado por etapa",
    "Framework KPI + roadmap 90 días",
    "Objetivo SQL/propuesta >= 55%"
])

add_subheading("Fase 3 — Escalado")

add_bullets([
    "Publicar casos de impacto",
    "Activar co-marketing sectorial",
    "Construir autoridad técnica"
])

# =============================================================================

add_chapter("6. KPIs RECOMENDADOS")

add_bullets([
    "Reuniones pre-agendadas",
    "SQL generados por feria",
    "Valor de pipeline",
    "Tiempo de ciclo comercial",
    "Ticket medio",
    "Ratio de estabilización contratada"
])

# =============================================================================

add_chapter("7. RECOMENDACIÓN EJECUTIVA FINAL")

add_body(
    "Para maximizar resultados en los próximos 12 meses, Ingecart debería "
    "concentrar recursos en cinco eventos core y utilizar eventos "
    "internacionales seleccionados como palancas de expansión y alianzas."
)

add_subheading("Eventos prioritarios")

add_bullets([
    "Corrugated Expo Barcelona",
    "FEFCO Summit Madrid",
    "Corrugated Week Fort Worth",
    "Empack Madrid",
    "Hispack Barcelona"
])

# =============================================================================
# FINAL STATEMENT
# =============================================================================

doc.add_paragraph("")

final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = final.add_run(
    "Engineering That Optimizes."
)

run.font.name = "Montserrat"
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = ORANGE

# =============================================================================
# SAVE
# =============================================================================

doc.save(OUTPUT_DOCX)

print("=================================================")
print("INGECART EXECUTIVE REPORT GENERATED")
print("=================================================")
print(f"Saved as: {OUTPUT_DOCX}")
print("=================================================")