import markdown2
from docx import Document
import os

# Paths
md_path = "informes/1 ESO/REPASO_COMPLETO_IMPRIMIBLE.md"
docx_path = "informes/1 ESO/REPASO_COMPLETO_IMPRIMIBLE.docx"

# Read markdown
with open(md_path, encoding="utf-8") as f:
    md_text = f.read()

# Convert markdown to HTML
html = markdown2.markdown(md_text)

# Create Word document
from bs4 import BeautifulSoup
soup = BeautifulSoup(html, "html.parser")
doc = Document()

def add_paragraph(text, style=None):
    doc.add_paragraph(text, style=style)

for elem in soup.find_all(['h1','h2','h3','h4','h5','h6','p','li']):
    if elem.name.startswith('h'):
        level = int(elem.name[1])
        doc.add_heading(elem.get_text(), level=level)
    elif elem.name == 'li':
        add_paragraph('• ' + elem.get_text())
    else:
        add_paragraph(elem.get_text())

# Save
os.makedirs(os.path.dirname(docx_path), exist_ok=True)
doc.save(docx_path)
print(f"Archivo Word generado: {docx_path}")
