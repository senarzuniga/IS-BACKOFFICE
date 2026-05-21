from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas


REPO_ROOT = Path(__file__).resolve().parents[1]
KIT_DIR = REPO_ROOT / "ingecart-marketing-kit" / "fespa-2026-kit-contenidos"
ASSETS_DIR = KIT_DIR / "assets"
PPTX_SOURCE = REPO_ROOT / "ingecart-marketing-kit" / "assets" / "presentations" / "Ingecart_FESPA_2026_Stand_Deck.pptx"

DEST_DIR = Path(r"C:\Users\Inaki Senar\Documents\INGECART\MARKETING\CONTENT\CONTENIDOS INGECART FESPA 2026")

LOGO_PATH = ASSETS_DIR / "ingeeniering.png"
HERO_PATH = ASSETS_DIR / "imagen_slogan_principal_ingecart.png"


def ensure_dest() -> None:
    DEST_DIR.mkdir(parents=True, exist_ok=True)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    if LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.5))

    if HERO_PATH.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic.add_run().add_picture(str(HERO_PATH), width=Inches(5.8))

    t = doc.add_paragraph()
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = t.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(5, 7, 11)

    st = doc.add_paragraph()
    st.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = st.add_run(subtitle)
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_page_break()


def is_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def parse_table_block(lines: List[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and is_table_line(lines[i]):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r"[-: ]+", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_to_doc(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if is_table_line(stripped):
            rows, nxt = parse_table_block(lines, i)
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=1, cols=cols)
                table.style = "Light List Accent 1"
                hdr = table.rows[0].cells
                for cidx, val in enumerate(rows[0]):
                    hdr[cidx].text = val
                for r in rows[1:]:
                    rr = table.add_row().cells
                    for cidx, val in enumerate(r):
                        rr[cidx].text = val
            i = nxt
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(stripped.lstrip("> "), style="Intense Quote")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", stripped), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+[\.)]\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+[\.)]\s+", "", stripped), style="List Number")
            i += 1
            continue

        p = doc.add_paragraph(stripped)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        i += 1


def build_docx_from_markdown(source_file: Path, output_file: Path, title: str, subtitle: str) -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, title, subtitle)
    add_markdown_to_doc(doc, source_file.read_text(encoding="utf-8"))
    doc.save(str(output_file))


def build_docx_from_text(source_file: Path, output_file: Path, title: str, subtitle: str) -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, title, subtitle)
    for line in source_file.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if not s:
            doc.add_paragraph("")
        elif s.endswith(":") or s.isupper():
            h = doc.add_paragraph(s)
            h.runs[0].bold = True
            h.runs[0].font.size = Pt(12)
        elif re.match(r"^\d+\)", s) or s.startswith("-"):
            cleaned = s.lstrip("- ")
            if re.match(r"^\d+\)", cleaned):
                cleaned = re.sub(r"^\d+\)\s*", "", cleaned)
                doc.add_paragraph(cleaned, style="List Number")
            else:
                doc.add_paragraph(cleaned, style="List Bullet")
        else:
            doc.add_paragraph(s)
    doc.save(str(output_file))


def draw_brand_header(c: canvas.Canvas, title: str, subtitle: str) -> None:
    w, h = A4
    c.setFillColor(colors.HexColor("#05070B"))
    c.rect(0, h - 42 * mm, w, 42 * mm, fill=1, stroke=0)

    c.setFillColor(colors.HexColor("#F4F5F7"))
    c.setFont("Helvetica-Bold", 18)
    c.drawString(18 * mm, h - 18 * mm, title)

    c.setFont("Helvetica", 11)
    c.setFillColor(colors.HexColor("#D0D4DA"))
    c.drawString(18 * mm, h - 27 * mm, subtitle)

    c.setFillColor(colors.HexColor("#FF6A00"))
    c.rect(18 * mm, h - 31 * mm, 30 * mm, 1.5 * mm, fill=1, stroke=0)

    if LOGO_PATH.exists():
        c.drawImage(str(LOGO_PATH), w - 48 * mm, h - 30 * mm, width=30 * mm, preserveAspectRatio=True, mask="auto")


def build_brochure_pdf(output_pdf: Path) -> None:
    c = canvas.Canvas(str(output_pdf), pagesize=A4)
    w, h = A4

    draw_brand_header(c, "Ingecart | FESPA 2026", "Brochure Comercial")

    y = h - 52 * mm
    if HERO_PATH.exists():
        c.drawImage(str(HERO_PATH), 18 * mm, y - 58 * mm, width=85 * mm, height=55 * mm, preserveAspectRatio=True, mask="auto")

    c.setFillColor(colors.HexColor("#05070B"))
    c.setFont("Helvetica-Bold", 15)
    c.drawString(108 * mm, y - 8 * mm, "Menos friccion operativa")
    c.drawString(108 * mm, y - 15 * mm, "Mas produccion util")

    c.setFillColor(colors.HexColor("#333333"))
    c.setFont("Helvetica", 10.5)
    txt = c.beginText(108 * mm, y - 24 * mm)
    paragraphs = [
        "Ingecart integra automatizacion, intralogistica e inteligencia operativa para plantas de corrugado.",
        "Soluciones destacadas: Paletizer + Easy Pack, Sistema Retal SR1400, Ingetrans 2800, AMR e Ing_PRO.",
        "Valor diferencial: ingenieria independiente, implantacion por fases y KPI medibles.",
        "Cifras: 28 anos de experiencia, 1.268 proyectos, 26 acuerdos internacionales, 194 instalaciones activas.",
    ]
    for p in paragraphs:
        txt.textLine(p)
        txt.textLine("")
    c.drawText(txt)

    c.setFillColor(colors.HexColor("#05070B"))
    c.setFont("Helvetica-Bold", 11)
    c.drawString(18 * mm, 28 * mm, "Contacto")
    c.setFont("Helvetica", 10)
    c.drawString(18 * mm, 22 * mm, "www.ingecart.eu | www.ingecart.eu/book-online | hablemos@ingecart.eu | +34 938 183 316")

    c.setFillColor(colors.HexColor("#FF6A00"))
    c.rect(0, 0, w, 8 * mm, fill=1, stroke=0)
    c.save()


def build_cards_pdf(output_pdf: Path) -> None:
    c = canvas.Canvas(str(output_pdf), pagesize=A4)
    w, h = A4
    draw_brand_header(c, "Tarjetas de Objeciones y Puentes", "Uso rapido de equipo de stand")

    cards = [
        ("Es muy caro", "Antes de precio, midamos coste de no actuar: parada, merma y mano de obra improductiva."),
        ("No quiero parar", "Implantacion por fases y ventanas de mantenimiento para minimizar impacto."),
        ("Ya tenemos proveedor", "Comparamos sobre el mismo baseline y decidimos por KPI, no por catalogo."),
        ("No estamos listos para IA", "Empezamos por caso concreto de alto impacto y adopcion rapida."),
        ("Mi layout es complejo", "Precisamente ahi aportamos mas: customizacion sobre flujo real."),
        ("No tengo CAPEX", "Definimos roadmap por etapas con hitos de retorno."),
    ]

    card_w = 86 * mm
    card_h = 48 * mm
    x_positions = [18 * mm, 108 * mm]
    y0 = h - 56 * mm

    for idx, (obj, rsp) in enumerate(cards):
        col = idx % 2
        row = idx // 2
        x = x_positions[col]
        y = y0 - row * (card_h + 8 * mm) - card_h

        c.setFillColor(colors.HexColor("#F8F8F8"))
        c.roundRect(x, y, card_w, card_h, 3 * mm, fill=1, stroke=0)
        c.setStrokeColor(colors.HexColor("#D0D0D0"))
        c.roundRect(x, y, card_w, card_h, 3 * mm, fill=0, stroke=1)

        c.setFillColor(colors.HexColor("#FF6A00"))
        c.rect(x, y + card_h - 8 * mm, card_w, 8 * mm, fill=1, stroke=0)

        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 9.5)
        c.drawString(x + 3 * mm, y + card_h - 5.7 * mm, f"Objecion: {obj}")

        c.setFillColor(colors.HexColor("#222222"))
        c.setFont("Helvetica", 9.5)
        text = c.beginText(x + 3 * mm, y + card_h - 14 * mm)
        for chunk in wrap_text(rsp, 45):
            text.textLine(chunk)
        c.drawText(text)

    c.save()


def wrap_text(text: str, max_chars: int) -> List[str]:
    words = text.split()
    lines: List[str] = []
    current = ""
    for w in words:
        candidate = (current + " " + w).strip()
        if len(candidate) <= max_chars:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines


def main() -> None:
    ensure_dest()

    mapping = [
        (
            KIT_DIR / "00_INFORME_COMPLETO_KIT_FERIA_FESPA_2026.md",
            DEST_DIR / "01_Informe_Maestro_FESPA_2026_Ingecart.docx",
            "Informe Maestro - Kit de Contenidos FESPA 2026",
            "Ingecart | Version final para uso comercial",
            "md",
        ),
        (
            KIT_DIR / "01_FICHAS_PRODUCTO_SERVICIO_FERIA.md",
            DEST_DIR / "02_Fichas_Producto_Servicio_FESPA_2026_Ingecart.docx",
            "Fichas de Producto y Servicio",
            "Ingecart | Uso comercial en stand",
            "md",
        ),
        (
            KIT_DIR / "02_BROCHURE_FERIA_FESPA_2026.md",
            DEST_DIR / "03_Brochure_FESPA_2026_Ingecart.docx",
            "Brochure Comercial FESPA 2026",
            "Ingecart | Material listo para compartir",
            "md",
        ),
        (
            KIT_DIR / "03_PRESENTACION_STAND_FESPA_2026.md",
            DEST_DIR / "04_Guion_Presentacion_Stand_FESPA_2026.docx",
            "Guion de Presentacion de Stand",
            "Ingecart | Script comercial-tecnico",
            "md",
        ),
        (
            KIT_DIR / "04_TARJETAS_OBJECIONES_Y_PUENTES_STAND.txt",
            DEST_DIR / "05_Tarjetas_Objeciones_Puentes_Stand_FESPA_2026.docx",
            "Tarjetas de Objeciones y Frases Puente",
            "Ingecart | Uso rapido de equipo",
            "txt",
        ),
        (
            KIT_DIR / "05_GUIA_APRENDIZAJE_ACELERADO_STAND.md",
            DEST_DIR / "06_Guia_Aprendizaje_Acelerado_Stand_FESPA_2026.docx",
            "Guia de Aprendizaje Acelerado",
            "Ingecart | Entrenamiento de equipo",
            "md",
        ),
    ]

    for src, out, title, subtitle, kind in mapping:
        if kind == "md":
            build_docx_from_markdown(src, out, title, subtitle)
        else:
            build_docx_from_text(src, out, title, subtitle)

    # PDF deliverables
    build_brochure_pdf(DEST_DIR / "03_Brochure_FESPA_2026_Ingecart.pdf")
    build_cards_pdf(DEST_DIR / "05_Tarjetas_Objeciones_Puentes_Stand_FESPA_2026.pdf")

    # Copy deck and assets
    if PPTX_SOURCE.exists():
        shutil.copy2(PPTX_SOURCE, DEST_DIR / "07_Ingecart_FESPA_2026_Stand_Deck.pptx")

    if LOGO_PATH.exists():
        shutil.copy2(LOGO_PATH, DEST_DIR / LOGO_PATH.name)

    if HERO_PATH.exists():
        shutil.copy2(HERO_PATH, DEST_DIR / HERO_PATH.name)

    # Index file
    index = DEST_DIR / "00_INDICE_ENTREGABLES_FESPA_2026.txt"
    index.write_text(
        "ENTREGABLES FINALES - CONTENIDOS INGECART FESPA 2026\n"
        "==============================================\n\n"
        "01_Informe_Maestro_FESPA_2026_Ingecart.docx\n"
        "02_Fichas_Producto_Servicio_FESPA_2026_Ingecart.docx\n"
        "03_Brochure_FESPA_2026_Ingecart.docx\n"
        "03_Brochure_FESPA_2026_Ingecart.pdf\n"
        "04_Guion_Presentacion_Stand_FESPA_2026.docx\n"
        "05_Tarjetas_Objeciones_Puentes_Stand_FESPA_2026.docx\n"
        "05_Tarjetas_Objeciones_Puentes_Stand_FESPA_2026.pdf\n"
        "06_Guia_Aprendizaje_Acelerado_Stand_FESPA_2026.docx\n"
        "07_Ingecart_FESPA_2026_Stand_Deck.pptx\n"
        "ingeeniering.png\n"
        "imagen_slogan_principal_ingecart.png\n",
        encoding="utf-8",
    )

    print(f"Entregables generados en: {DEST_DIR}")


if __name__ == "__main__":
    main()
