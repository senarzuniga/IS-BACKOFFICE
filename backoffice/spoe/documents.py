from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import Dict, List

from docx import Document

from .models import OfferInput


_DOC_SECTIONS = [
    "Commercial Offer",
    "Technical Proposal",
    "Executive Summary",
    "Bill of Materials",
    "Scope of Supply",
    "Excluded Scope",
    "Installation Estimate",
    "Commissioning",
    "Commercial Conditions",
    "General Terms",
    "Engineering Annex",
]


def _lang(offer: OfferInput, es: str, en: str) -> str:
    return es if offer.language.lower().startswith("es") else en


def _add_offer_header(doc: Document, offer: OfferInput) -> None:
    doc.add_heading(f"SPOE | SR1400 | {offer.offer_number}", 0)
    doc.add_paragraph(f"Customer: {offer.customer}")
    doc.add_paragraph(f"Plant: {offer.plant} | Country: {offer.country}")
    doc.add_paragraph(f"Project: {offer.project_name}")
    doc.add_paragraph(f"Date: {offer.offer_date.isoformat()}")


def _build_doc(section: str, offer: OfferInput, bom: Dict[str, int], knowledge: Dict) -> Document:
    doc = Document()
    _add_offer_header(doc, offer)
    doc.add_heading(section, level=1)

    if section == "Bill of Materials":
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Component"
        table.rows[0].cells[1].text = "Quantity"
        for k, v in bom.items():
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)
    elif section == "Executive Summary":
        doc.add_paragraph(_lang(offer, knowledge["executive"]["es"], knowledge["executive"]["en"]))
    elif section == "Technical Proposal":
        doc.add_paragraph(_lang(offer, knowledge["technical_description"]["es"], knowledge["technical_description"]["en"]))
        doc.add_paragraph(_lang(offer, knowledge["design_philosophy"]["es"], knowledge["design_philosophy"]["en"]))
    else:
        body = _lang(
            offer,
            f"Seccion {section} generada automaticamente por SPOE para SR1400.",
            f"Section {section} generated automatically by SPOE for SR1400.",
        )
        doc.add_paragraph(body)

    if offer.layout_image_path:
        doc.add_paragraph(
            _lang(
                offer,
                f"Layout referenciado: {offer.layout_image_path}",
                f"Layout referenced: {offer.layout_image_path}",
            )
        )
    if offer.optional_attachment_paths:
        doc.add_paragraph(
            _lang(
                offer,
                "Adjuntos incluidos:",
                "Included attachments:",
            )
        )
        for att in offer.optional_attachment_paths:
            doc.add_paragraph(f"- {att}")

    return doc


def generate_offer_documents(offer: OfferInput, bom: Dict[str, int], knowledge: Dict) -> Dict[str, str]:
    output_dir = Path("reports/spoe/generated") / offer.offer_number
    output_dir.mkdir(parents=True, exist_ok=True)

    generated: Dict[str, str] = {}
    for section in _DOC_SECTIONS:
        doc = _build_doc(section, offer, bom, knowledge)
        filename = section.lower().replace(" ", "_") + ".docx"
        file_path = output_dir / filename
        doc.save(file_path)
        generated[section] = str(file_path)

    manifest = output_dir / "manifest.txt"
    lines = [f"Generated at: {datetime.now(UTC).isoformat()}"]
    lines.extend([f"{k}: {v}" for k, v in generated.items()])
    manifest.write_text("\n".join(lines), encoding="utf-8")
    return generated
