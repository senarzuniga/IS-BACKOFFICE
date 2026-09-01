from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .models import Proposal


REPO_ROOT = Path(__file__).resolve().parent.parent.parent
LOGO_CANDIDATES = [
    REPO_ROOT / "assets" / "branding" / "ingecart_logo.png",
    REPO_ROOT / "scrapes" / "ingecart_proyectos" / "assets" / "Copia-de-ingeeniering.png",
]

ORANGE = "F36B21"
DARK = "171717"
WHITE = "FFFFFF"
LIGHT_GREY = "F7F7F7"


class ProposalWordGenerator:
    """Build Word offers in Ingecart style derived from the PAIGE offer model."""

    def generate(self, proposal: Proposal, output_dir: Path) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"{(proposal.number or proposal.id or 'proposal').replace('/', '-')}.docx"
        out_path = output_dir / filename

        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        self._configure_styles(doc)
        self._configure_header(doc, proposal)
        self._add_cover(doc, proposal)
        self._add_services_table(doc, proposal)
        self._add_rich_sections(doc, proposal)
        doc.save(out_path)
        return out_path

    def _configure_styles(self, doc: Document) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = RGBColor.from_string(DARK)
        normal.paragraph_format.space_after = Pt(5)
        for style_name, size, color in (("Heading 1", 17, ORANGE), ("Heading 2", 12, DARK)):
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)

    def _configure_header(self, doc: Document, proposal: Proposal) -> None:
        header = doc.sections[0].header
        table = header.add_table(rows=1, cols=2, width=Inches(6.7))
        table.autofit = False
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        logo = self._pick_logo()
        if logo:
            left.paragraphs[0].add_run().add_picture(str(logo), width=Inches(1.1))
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta = right.paragraphs[0].add_run(f"{proposal.number or 'OFF-XXXX'}\n{datetime.now():%Y-%m-%d}")
        meta.font.name = "Arial"
        meta.font.size = Pt(8)
        meta.font.bold = True

    def _add_cover(self, doc: Document, proposal: Proposal) -> None:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pr = title._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), DARK)
        p_pr.append(shd)
        run = title.add_run(
            f"\nOFERTA ANUAL DE MANTENIMIENTO\n{proposal.customer.upper() if proposal.customer else 'CLIENTE'}\n"
        )
        run.font.name = "Arial"
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2_pr = subtitle._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:fill"), ORANGE)
        p2_pr.append(shd2)
        run2 = subtitle.add_run("INGEPRO MONITORIZACION · AI PREDICTIVA · CANAL INDUSTRIAL DE RECAMBIOS")
        run2.font.name = "Arial"
        run2.font.size = Pt(10)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor.from_string(WHITE)

        info = doc.add_table(rows=5, cols=2)
        info.autofit = False
        rows = (
            ("Cliente", proposal.customer or "-"),
            ("Planta", proposal.plant or "-"),
            ("Referencia", proposal.number or "-"),
            ("Fecha", proposal.date_created[:10] if proposal.date_created else datetime.now().strftime("%Y-%m-%d")),
            ("Duracion", proposal.duration or "12 meses"),
        )
        for idx, (label, value) in enumerate(rows):
            left = info.rows[idx].cells[0]
            right = info.rows[idx].cells[1]
            left.text = label
            right.text = value
            self._shade_cell(left, DARK)
            self._shade_cell(right, LIGHT_GREY)
            self._format_cell(left, bold=True, color=WHITE)
            self._format_cell(right, bold=idx == 0)

    def _add_services_table(self, doc: Document, proposal: Proposal) -> None:
        doc.add_paragraph()
        doc.add_heading("Servicios anuales incluidos", level=1)
        table = doc.add_table(rows=1, cols=4)
        headers = ("Servicio", "Frecuencia", "Cobertura", "Importe anual")
        for idx, name in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = name
            self._shade_cell(cell, DARK)
            self._format_cell(cell, bold=True, color=WHITE)
        for service in proposal.services:
            if not service.enabled:
                continue
            row = table.add_row().cells
            row[0].text = service.name
            row[1].text = service.frequency or service.unit or "-"
            row[2].text = service.coverage or "-"
            row[3].text = f"{service.total_price:,.0f} {proposal.currency}"
            for cell in row:
                self._format_cell(cell)
        total = proposal.total_price + proposal.optional_price
        summary = table.add_row().cells
        summary[0].text = "TOTAL ANUAL RECOMENDADO"
        summary[0].merge(summary[2])
        summary[3].text = f"{total:,.0f} {proposal.currency}"
        self._shade_cell(summary[0], ORANGE)
        self._shade_cell(summary[3], ORANGE)
        self._format_cell(summary[0], bold=True, color=WHITE)
        self._format_cell(summary[3], bold=True, color=WHITE)

    def _add_rich_sections(self, doc: Document, proposal: Proposal) -> None:
        sections = [
            ("Resumen ejecutivo", proposal.executive_summary),
            ("Programa de mantenimiento", proposal.maintenance_programme),
            ("INGEPRO monitorizacion y AI", proposal.ingpro_section),
            ("Entregables", proposal.deliverables),
            ("Condiciones comerciales", proposal.commercial_conditions),
            ("Resumen economico", proposal.pricing_summary),
            ("Aceptacion", proposal.acceptance),
        ]
        for title, content in sections:
            normalized = self._html_to_text(content)
            if not normalized:
                continue
            doc.add_heading(title, level=1)
            for block in [line.strip() for line in normalized.splitlines() if line.strip()]:
                doc.add_paragraph(block)

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)
        shd.set(qn("w:val"), "clear")

    @staticmethod
    def _format_cell(cell, *, bold: bool = False, color: str = DARK) -> None:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8.5)
                run.font.bold = bold
                run.font.color.rgb = RGBColor.from_string(color)

    def _pick_logo(self) -> Path | None:
        for candidate in LOGO_CANDIDATES:
            if candidate.exists():
                return candidate
        return None

    @staticmethod
    def _html_to_text(value: str) -> str:
        content = (value or "").strip()
        if not content:
            return ""
        soup = BeautifulSoup(content, "html.parser")
        text = soup.get_text("\n", strip=True)
        text = re.sub(r"\n{2,}", "\n", text)
        return text.strip()
