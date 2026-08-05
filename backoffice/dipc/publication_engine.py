from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from odf.opendocument import OpenDocumentText
from odf.text import H, P
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .component_library import render_component
from .models import DocumentModel
from .theme_engine import build_css


class PublicationEngine:
    def export_all(self, document: DocumentModel, output_dir: str | Path, corporate_css: str | None = None) -> dict[str, str]:
        root = Path(output_dir)
        root.mkdir(parents=True, exist_ok=True)
        css_path = root / "dipc_theme.css"
        if corporate_css is not None and corporate_css.strip():
            css_path.write_text(corporate_css, encoding="utf-8")
        else:
            css_path.write_text(build_css(document.theme_variant, document.metadata.get("palette")), encoding="utf-8")

        html_path = root / "index.html"
        markdown_path = root / "document.md"
        docx_path = root / "document.docx"
        pdf_path = root / "document.pdf"
        odt_path = root / "document.odt"
        presentation_html_path = root / "presentation.html"
        microsite_dir = root / "microsite"
        portal_dir = root / "portal_web"
        json_path = root / "document_model.json"

        html = self.render_html(document, css_href="dipc_theme.css")
        html_path.write_text(html, encoding="utf-8")
        markdown_path.write_text(self.render_markdown(document), encoding="utf-8")
        json_path.write_text(document.model_dump_json(indent=2), encoding="utf-8")
        self.render_docx(document, docx_path)
        self.render_pdf(document, pdf_path)
        self.render_odt(document, odt_path)
        presentation_html_path.write_text(self.render_presentation_html(document, css_href="dipc_theme.css"), encoding="utf-8")

        self._build_site(root, microsite_dir, html_path, css_path, title=document.title)
        self._build_site(root, portal_dir, html_path, css_path, title=document.title)

        return {
            "html": str(html_path),
            "markdown": str(markdown_path),
            "docx": str(docx_path),
            "pdf": str(pdf_path),
            "odt": str(odt_path),
            "presentation_html": str(presentation_html_path),
            "portal_web": str(portal_dir / "index.html"),
            "microsite": str(microsite_dir / "index.html"),
            "document_json": str(json_path),
            "theme_css": str(css_path),
        }

    def render_html(self, document: DocumentModel, css_href: str) -> str:
        translations = document.metadata.get("translations", {}) if isinstance(document.metadata, dict) else {}
        en = translations.get("en", {}) if isinstance(translations, dict) else {}
        es = translations.get("es", {}) if isinstance(translations, dict) else {}
        title_en = en.get("title", document.title)
        title_es = es.get("title", document.title)
        subtitle_en = en.get("subtitle", document.subtitle or "")
        subtitle_es = es.get("subtitle", document.subtitle or "")

        nav = []
        body = []
        for section in document.sections:
            nav.append(f"<a href='#section-{section.order}'>{section.order}. {section.title}</a>")
            rendered = []
            for block in section.blocks:
                for component in block.components:
                    rendered.append(render_component(component))
            body.append(
                f"<section class='dipc-section' id='section-{section.order}'><h2>{section.title}</h2><p>{section.summary or ''}</p>{''.join(rendered)}</section>"
            )

        return f"""<!doctype html>
<html lang='en'>
<head>
  <meta charset='utf-8' />
  <meta name='viewport' content='width=device-width, initial-scale=1' />
  <title>{document.title}</title>
  <meta name='description' content='{document.subtitle or document.title}' />
  <link rel='stylesheet' href='{css_href}' />
</head>
<body>
    <div class='dipc-lang-switch' role='group' aria-label='Language selector'>
        <span>Language</span>
        <button type='button' data-lang='en'>English</button>
        <button type='button' data-lang='es'>Español</button>
    </div>
  <div class='dipc-shell'>
    <aside class='dipc-sidebar'>
            <h2 data-i18n-title data-en="{title_en}" data-es="{title_es}">{document.title}</h2>
            <p data-i18n-subtitle data-en="{subtitle_en}" data-es="{subtitle_es}">{document.subtitle or ''}</p>
      {''.join(nav)}
    </aside>
    <main class='dipc-main'>
            <header class='dipc-hero'>
                <h1 data-i18n-title data-en="{title_en}" data-es="{title_es}">{document.title}</h1>
                <p data-i18n-subtitle data-en="{subtitle_en}" data-es="{subtitle_es}">{document.subtitle or ''}</p>
            </header>
      {''.join(body)}
    </main>
  </div>
    <script>
        (function() {{
            const applyLang = (lang) => {{
                document.querySelectorAll('[data-i18n-title]').forEach((el) => {{
                    const next = el.getAttribute('data-' + lang);
                    if (next) el.textContent = next;
                }});
                document.querySelectorAll('[data-i18n-subtitle]').forEach((el) => {{
                    const next = el.getAttribute('data-' + lang);
                    if (next) el.textContent = next;
                }});
                document.documentElement.setAttribute('lang', lang);
            }};

            document.querySelectorAll('.dipc-lang-switch button[data-lang]').forEach((btn) => {{
                btn.addEventListener('click', () => applyLang(btn.getAttribute('data-lang') || 'en'));
            }});

            applyLang('en');
        }})();
    </script>
</body>
</html>"""

    def render_presentation_html(self, document: DocumentModel, css_href: str) -> str:
        slides = []
        for section in document.sections:
            components = []
            for block in section.blocks:
                for component in block.components:
                    components.append(render_component(component))
            slides.append(f"<section class='dipc-section slide'><h2>{section.title}</h2>{''.join(components)}</section>")
        return f"""<!doctype html>
<html lang='en'>
<head><meta charset='utf-8' /><meta name='viewport' content='width=device-width, initial-scale=1' /><title>{document.title} - Presentation</title><link rel='stylesheet' href='{css_href}' /></head>
<body><main>{''.join(slides)}</main></body></html>"""

    def render_markdown(self, document: DocumentModel) -> str:
        lines = [f"# {document.title}", "", document.subtitle or "", ""]
        for section in document.sections:
            lines.append(f"## {section.order}. {section.title}")
            if section.summary:
                lines.append(section.summary)
            lines.append("")
            for block in section.blocks:
                for component in block.components:
                    if component.title:
                        lines.append(f"### {component.title}")
                    if component.body:
                        lines.append(component.body)
                    for item in component.items:
                        title = item.get("title")
                        body = item.get("body")
                        if title or body:
                            lines.append(f"- {title or ''}: {body or ''}".strip())
                    rows = component.props.get("rows", [])
                    if rows:
                        for row in rows:
                            lines.append(" | ".join(str(cell) for cell in row))
                    lines.append("")
        return "\n".join(lines)

    def render_docx(self, document: DocumentModel, path: str | Path) -> None:
        doc = DocxDocument()
        doc.add_heading(document.title, level=0)
        if document.subtitle:
            doc.add_paragraph(document.subtitle)
        for section in document.sections:
            doc.add_heading(f"{section.order}. {section.title}", level=1)
            if section.summary:
                doc.add_paragraph(section.summary)
            for block in section.blocks:
                for component in block.components:
                    if component.title:
                        doc.add_heading(component.title, level=2)
                    if component.body:
                        doc.add_paragraph(component.body)
                    for item in component.items:
                        doc.add_paragraph(f"{item.get('title', '')}: {item.get('body', '')}", style="List Bullet")
                    rows = component.props.get("rows", [])
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        for ridx, row in enumerate(rows):
                            for cidx, cell in enumerate(row):
                                table.rows[ridx].cells[cidx].text = str(cell)
        doc.save(str(path))

    def render_pdf(self, document: DocumentModel, path: str | Path) -> None:
        styles = getSampleStyleSheet()
        story: list[Any] = [Paragraph(document.title, styles["Title"])]
        if document.subtitle:
            story.extend([Paragraph(document.subtitle, styles["BodyText"]), Spacer(1, 12)])
        for section in document.sections:
            story.append(Paragraph(f"{section.order}. {section.title}", styles["Heading1"]))
            if section.summary:
                story.append(Paragraph(section.summary, styles["BodyText"]))
            for block in section.blocks:
                for component in block.components:
                    if component.title:
                        story.append(Paragraph(component.title, styles["Heading2"]))
                    if component.body:
                        story.append(Paragraph(component.body, styles["BodyText"]))
                    for item in component.items:
                        story.append(Paragraph(f"- {item.get('title', '')}: {item.get('body', '')}", styles["BodyText"]))
                    story.append(Spacer(1, 8))
        SimpleDocTemplate(str(path), pagesize=A4).build(story)

    def render_odt(self, document: DocumentModel, path: str | Path) -> None:
        odt = OpenDocumentText()
        odt.text.addElement(H(outlinelevel=1, text=document.title))
        if document.subtitle:
            odt.text.addElement(P(text=document.subtitle))
        for section in document.sections:
            odt.text.addElement(H(outlinelevel=2, text=f"{section.order}. {section.title}"))
            if section.summary:
                odt.text.addElement(P(text=section.summary))
            for block in section.blocks:
                for component in block.components:
                    if component.title:
                        odt.text.addElement(H(outlinelevel=3, text=component.title))
                    if component.body:
                        odt.text.addElement(P(text=component.body))
                    for item in component.items:
                        odt.text.addElement(P(text=f"{item.get('title', '')}: {item.get('body', '')}"))
        odt.save(str(path), addsuffix=False)

    def _build_site(self, root: Path, site_dir: Path, html_path: Path, css_path: Path, title: str) -> None:
        site_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(html_path, site_dir / "index.html")
        shutil.copy2(css_path, site_dir / "dipc_theme.css")
        manifest = {
            "title": title,
            "index": "index.html",
            "theme": "dipc_theme.css",
            "generated_from": str(root),
        }
        (site_dir / "site_manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")
